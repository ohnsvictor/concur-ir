from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy, datetime

# ─── HELPERS ─────────────────────────────────────────────────────────────────

def set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

def set_cell_borders(cell, top=None, bottom=None, left=None, right=None):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for side, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        if val:
            el = OxmlElement(f'w:{side}')
            el.set(qn('w:val'), val.get('val', 'single'))
            el.set(qn('w:sz'), str(val.get('sz', 4)))
            el.set(qn('w:color'), val.get('color', 'auto'))
            tcBorders.append(el)
    tcPr.append(tcBorders)

def no_space_before_after(para):
    pPr = para._p.get_or_add_pPr()
    spacing = OxmlElement('w:spacing')
    spacing.set(qn('w:before'), '0')
    spacing.set(qn('w:after'), '0')
    pPr.append(spacing)

def add_run(para, text, bold=False, italic=False, size=None, color=None, font_name='Calibri'):
    run = para.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.name = font_name
    if size:
        run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor(*color)
    return run

def merge_row(table, row_idx):
    row = table.rows[row_idx]
    row.cells[0].merge(row.cells[1])
    row.cells[0].merge(row.cells[2])
    return row.cells[0]

# ─── COLORS ──────────────────────────────────────────────────────────────────

C_DOMAIN_BG    = '1C3A6E'   # dark navy — domain headers
C_DOMAIN_TXT   = 'FFFFFF'   # white text
C_COND_HEADER  = '2E75B6'   # medium blue — conditional design label
C_COND_ITEM    = 'EEF3FA'   # very light blue — conditional items
C_NOTE_BG      = 'F5F7FA'   # near-white — note text in D2
C_SUB_BG       = 'FAFAFA'   # off-white — → sub-items
C_PAGE_HEADER  = 'E8EDF5'   # light blue-gray — column header row
C_WHITE        = 'FFFFFF'
C_SECTION_SEP  = 'D0D8E8'   # light line between domains

# ─── DATA ────────────────────────────────────────────────────────────────────

DOMAINS = [

    # ── D1 ────────────────────────────────────────────────────────────────────
    {
        "id": "D1", "title": "STUDY DESIGN & PRE-REGISTRATION",
        "items": [
            {"id": "1.1",  "text": "Study design explicitly stated",
             "sub": ["Prospective vs. retrospective", "Observational vs. interventional", "Single-center vs. multi-center"],
             "rec": False},
            {"id": "1.2",  "text": "Study period reported (enrollment start and end dates)", "sub": [], "rec": False},
            {"id": "1.3",  "text": "Ethics committee approval stated with institutional reference number", "sub": [], "rec": False},
            {"id": "1.4",  "text": "Informed consent process described",
             "sub": ["If individual consent: confirmation provided",
                     "If waived: justification and IRB acknowledgment stated",
                     "If registry-based: data use agreement described"],
             "rec": False},
            {"id": "1.5",  "text": "Prospective pre-registration declared (ClinicalTrials.gov, UMIN, ISRCTN, ReBEC, etc.)", "sub": [], "rec": True},
        ],
        "conditionals": [
            {"design": "RCT", "note": "", "items": [
                {"id": "1.RCT.1", "text": "Trial registration number reported with primary endpoint as registered", "sub": [], "rec": False},
                {"id": "1.RCT.2", "text": "Post-registration protocol amendments disclosed with rationale", "sub": [], "rec": True},
            ]},
            {"design": "Registry", "note": "", "items": [
                {"id": "1.REG.1", "text": "Registry governance and data custodian identified", "sub": [], "rec": False},
                {"id": "1.REG.2", "text": "Registry start date vs. study extraction period distinguished", "sub": [], "rec": False},
            ]},
            {"design": "Case Series", "note": "", "items": [
                {"id": "1.CS.1", "text": "IDEAL developmental stage declared (Stage 1, 2a, 2b, or 3)", "sub": [], "rec": True},
                {"id": "1.CS.2", "text": "Whether technique was novel or established at time of study stated", "sub": [], "rec": False},
            ]},
        ]
    },

    # ── D2 ────────────────────────────────────────────────────────────────────
    {
        "id": "D2", "title": "PATIENT SELECTION & ELIGIBILITY",
        "items": [
            {"id": "2.1", "text": "Inclusion criteria explicitly stated", "sub": [], "rec": False},
            {"id": "2.2", "text": "Exclusion criteria explicitly stated", "sub": [], "rec": False},
            {"id": "2.3", "text": "Patient selection process described (consecutive, random, or convenience — with rationale)", "sub": [], "rec": False},
            {"id": "2.4", "text": "Patient flow diagram provided (screened → eligible → enrolled → analyzed)", "sub": [], "rec": True},
        ],
        "conditionals": [
            {"design": "Case Series",
             "note": "In case series, selection bias is the primary threat to validity. Authors should clearly state whether all eligible patients were included and whether enrollment was consecutive.",
             "items": [
                {"id": "2.CS.1", "text": "Consecutive enrollment confirmed, or deviation explicitly acknowledged", "sub": [], "rec": True},
                {"id": "2.CS.2", "text": "Whether all eligible patients were included or only a subset, and why", "sub": [], "rec": False},
            ]},
            {"design": "Cohort",
             "note": "Cohort studies in IR frequently draw from institutional databases. The source population and any filtering from the broader registry to the study cohort must be transparent to allow assessment of generalizability.",
             "items": [
                {"id": "2.CO.1", "text": "Source population clearly defined (institutional registry, referral base, or community-based)", "sub": [], "rec": False},
                {"id": "2.CO.2", "text": "Comparator group selection methodology described", "sub": [], "rec": True},
            ]},
            {"design": "RCT",
             "note": "RCTs in IR are prone to post-randomization exclusions due to technical failures or anatomical unsuitability discovered at the time of the procedure. These must be distinguished from pre-procedural exclusions.",
             "items": [
                {"id": "2.RCT.1", "text": "Screening log with reasons for exclusion reported", "sub": [], "rec": True},
                {"id": "2.RCT.2", "text": "Post-randomization exclusions (e.g. technical failure, anatomical unsuitability) distinguished from pre-procedural exclusions", "sub": [], "rec": False},
                {"id": "2.RCT.3", "text": "Run-in period described if applicable", "sub": [], "rec": False},
            ]},
            {"design": "Registry",
             "note": "Registry-based IR studies carry the risk of differential inclusion. Authors should quantify completeness and discuss potential enrollment bias.",
             "items": [
                {"id": "2.REG.1", "text": "Proportion of eligible registry patients included reported", "sub": [], "rec": False},
                {"id": "2.REG.2", "text": "Potential selection bias from registry enrollment criteria discussed", "sub": [], "rec": True},
            ]},
        ]
    },

    # ── D3a ───────────────────────────────────────────────────────────────────
    {
        "id": "D3a", "title": "BASELINE CLINICAL CHARACTERIZATION",
        "items": [
            {"id": "3a.1", "text": "Age and sex reported for all patients", "sub": [], "rec": False},
            {"id": "3a.2", "text": "Relevant comorbidities reported (at minimum: diabetes, hypertension, CKD, coagulopathy, prior oncological treatment when applicable)", "sub": [], "rec": False},
            {"id": "3a.3", "text": "Performance status or functional scale reported (e.g. ECOG, Karnofsky, ASA — appropriate to clinical context)", "sub": [], "rec": True},
            {"id": "3a.4", "text": "Relevant laboratory values reported (e.g. creatinine, INR, platelets, tumor markers — contextual to procedure)", "sub": [], "rec": False},
            {"id": "3a.5", "text": "Prior treatments related to current indication described (e.g. prior systemic therapy, locoregional treatment, surgical attempt)", "sub": [], "rec": False},
            {"id": "3a.6", "text": "Relevant current medications reported (anticoagulation, antiplatelet, immunosuppression — where procedurally relevant)", "sub": [], "rec": True},
        ],
        "conditionals": [
            {"design": "Cohort with Comparator / RCT", "note": "", "items": [
                {"id": "3a.C1", "text": "Baseline characteristics reported separately per group", "sub": [], "rec": False},
                {"id": "3a.C2", "text": "Clinically meaningful imbalances between groups identified and addressed in statistical analysis", "sub": [], "rec": True},
            ]},
            {"design": "RCT", "note": "", "items": [
                {"id": "3a.RCT.1", "text": "Standardized mean differences reported for key baseline variables (recommended — preferred over p-values for baseline comparisons)", "sub": [], "rec": True},
            ]},
            {"design": "Registry", "note": "", "items": [
                {"id": "3a.REG.1", "text": "Completeness of each baseline variable reported (% available)", "sub": [], "rec": False},
                {"id": "3a.REG.2", "text": "Missing baseline data handling described (imputation method or sensitivity analysis)", "sub": [], "rec": True},
            ]},
        ]
    },

    # ── D3b ───────────────────────────────────────────────────────────────────
    {
        "id": "D3b", "title": "BASELINE ANATOMICAL & RADIOLOGICAL CHARACTERIZATION",
        "items": [
            {"id": "3b.1", "text": "Imaging modality used for baseline assessment stated, with date relative to procedure", "sub": [], "rec": False},
            {"id": "3b.2", "text": "Target anatomy described with objective measurements (lesion size, vessel diameter, stenosis degree, thrombus extent — contextual to procedure)", "sub": [], "rec": False},
            {"id": "3b.3", "text": "Anatomical complexity reported (tortuosity, calcification, prior surgical anatomy, access challenges)", "sub": [], "rec": True},
            {"id": "3b.4", "text": "Validated staging or complexity classification applied when one exists (BCLC, TASC II, Rutherford, Child-Pugh, mRECIST, etc.)", "sub": [], "rec": True},
            {"id": "3b.5", "text": "Central vs. local imaging review stated", "sub": [], "rec": True},
        ],
        "conditionals": [
            {"design": "Cohort with Comparator / RCT", "note": "", "items": [
                {"id": "3b.C1", "text": "Anatomical characteristics reported separately per group", "sub": [], "rec": False},
                {"id": "3b.C2", "text": "Anatomical complexity included as covariate in adjusted analyses when groups differ", "sub": [], "rec": True},
            ]},
            {"design": "Registry", "note": "", "items": [
                {"id": "3b.REG.1", "text": "Prospectively vs. retrospectively collected anatomical variables distinguished", "sub": [], "rec": False},
            ]},
        ]
    },

    # ── D4 ────────────────────────────────────────────────────────────────────
    {
        "id": "D4", "title": "INTERVENTION INDICATION & ALTERNATIVES",
        "items": [
            {"id": "4.1", "text": "Clinical indication for the procedure explicitly stated", "sub": [], "rec": False},
            {"id": "4.2", "text": "Indication based on objective and reproducible criteria (imaging, laboratory, functional — not operator discretion alone)", "sub": [], "rec": True},
            {"id": "4.3", "text": "Alternative treatment options considered and described (surgical, medical, other IR approach)", "sub": [], "rec": True},
            {"id": "4.4", "text": "Reason for choosing IR over alternatives stated", "sub": [], "rec": True},
            {"id": "4.5", "text": "Whether IR was first-line, second-line, or rescue therapy explicitly declared", "sub": [], "rec": False},
        ],
        "conditionals": [
            {"design": "Case Series", "note": "", "items": [
                {"id": "4.CS.1", "text": "Whether indication was homogeneous across all patients or varied — stated explicitly", "sub": [], "rec": False},
            ]},
            {"design": "Cohort with Comparator / RCT", "note": "", "items": [
                {"id": "4.C1", "text": "Indication criteria confirmed as applied uniformly across both groups", "sub": [], "rec": False},
            ]},
            {"design": "RCT", "note": "", "items": [
                {"id": "4.RCT.1", "text": "Equipoise justification provided (why both arms were considered clinically acceptable at time of enrollment)", "sub": [], "rec": True},
            ]},
            {"design": "Registry", "note": "", "items": [
                {"id": "4.REG.1", "text": "Whether indication was protocol-driven or left to operator discretion — stated explicitly", "sub": [], "rec": False},
            ]},
        ]
    },

    # ── D5 ────────────────────────────────────────────────────────────────────
    {
        "id": "D5", "title": "OPERATOR & CENTER CHARACTERISTICS",
        "items": [
            {"id": "5.1", "text": "Number of operators involved reported", "sub": [], "rec": False},
            {"id": "5.2", "text": "Operator experience in the specific procedure stated (years of experience, number of prior cases, or formal training — at least one metric)", "sub": [], "rec": False},
            {"id": "5.3", "text": "Annual center volume for the specific procedure reported", "sub": [], "rec": True},
            {"id": "5.4", "text": "Academic vs. community center declared", "sub": [], "rec": False},
            {"id": "5.5", "text": "Whether the study period overlapped with a learning curve explicitly addressed", "sub": [], "rec": True},
        ],
        "conditionals": [
            {"design": "Case Series", "note": "", "items": [
                {"id": "5.CS.1", "text": "Single-operator vs. multi-operator series declared", "sub": [], "rec": False},
                {"id": "5.CS.2", "text": "If multi-operator: inter-operator variability in technique acknowledged", "sub": [], "rec": True},
            ]},
            {"design": "Cohort with Comparator / RCT", "note": "", "items": [
                {"id": "5.C1", "text": "Operator experience comparable between groups, or difference addressed in analysis", "sub": [], "rec": False},
            ]},
            {"design": "RCT", "note": "", "items": [
                {"id": "5.RCT.1", "text": "Operator credentialing criteria pre-specified (minimum experience required to participate)", "sub": [], "rec": False},
                {"id": "5.RCT.2", "text": "Operator volume during trial period reported", "sub": [], "rec": True},
            ]},
            {"design": "Registry", "note": "", "items": [
                {"id": "5.REG.1", "text": "Number of participating centers reported", "sub": [], "rec": False},
                {"id": "5.REG.2", "text": "Variability in operator experience and center volume across sites described", "sub": [], "rec": True},
            ]},
        ]
    },

    # ── D6a ───────────────────────────────────────────────────────────────────
    {
        "id": "D6a", "title": "PROCEDURE DESCRIPTION — TECHNICAL EXECUTION",
        "items": [
            {"id": "6a.1", "text": "Access route described (arterial, venous, percutaneous, transhepatic, etc.)", "sub": [], "rec": False},
            {"id": "6a.2", "text": "Step-by-step technical description provided with sufficient detail for reproducibility", "sub": [], "rec": False},
            {"id": "6a.3", "text": "Predefined procedural protocol or manual used — stated explicitly, or absence acknowledged", "sub": [], "rec": False},
            {"id": "6a.4", "text": "Allowed technical variations described", "sub": [], "rec": True},
            {"id": "6a.5", "text": "Technical success defined with objective, pre-specified criteria", "sub": [], "rec": False},
            {"id": "6a.6", "text": "Bailout strategies described", "sub": [], "rec": True},
            {"id": "6a.7", "text": "Technical failures reported separately from clinical failures", "sub": [], "rec": False},
        ],
        "conditionals": [
            {"design": "Case Series", "note": "", "items": [
                {"id": "6a.CS.1", "text": "Whether technique evolved during the study period acknowledged", "sub": [], "rec": True},
            ]},
            {"design": "Cohort with Comparator / RCT", "note": "", "items": [
                {"id": "6a.C1", "text": "Technical protocol confirmed as applied uniformly across both groups, or differences described", "sub": [], "rec": False},
            ]},
            {"design": "RCT", "note": "", "items": [
                {"id": "6a.RCT.1", "text": "Protocol deviations documented and reported", "sub": [], "rec": False},
                {"id": "6a.RCT.2", "text": "Co-interventions applied equally to both arms described", "sub": [], "rec": True},
            ]},
            {"design": "Registry", "note": "", "items": [
                {"id": "6a.REG.1", "text": "Degree of technical standardization across participating centers described", "sub": [], "rec": False},
            ]},
        ]
    },

    # ── D6b ───────────────────────────────────────────────────────────────────
    {
        "id": "D6b", "title": "PROCEDURE DESCRIPTION — IMAGING GUIDANCE",
        "items": [
            {"id": "6b.1", "text": "Imaging modality used for guidance stated (fluoroscopy, US, CT, CBCT, MRI, fusion, etc.)", "sub": [], "rec": False},
            {"id": "6b.2", "text": "Real-time guidance vs. pre-procedural planning only — distinguished explicitly", "sub": [], "rec": False},
            {"id": "6b.3", "text": "Contrast agent used intraprocedurally described (type, concentration, volume — when applicable)", "sub": [], "rec": True},
            {"id": "6b.4", "text": "Navigation or fusion software used identified (manufacturer, system name)", "sub": [], "rec": True},
            {"id": "6b.5", "text": "Image quality or technical limitations of guidance acknowledged when relevant", "sub": [], "rec": True},
        ],
        "conditionals": [
            {"design": "Case Series", "note": "", "items": [
                {"id": "6b.CS.1", "text": "Changes in imaging guidance modality during the study period disclosed", "sub": [], "rec": True},
            ]},
            {"design": "RCT", "note": "", "items": [
                {"id": "6b.RCT.1", "text": "Imaging guidance protocol standardized across sites and operators — confirmed or deviation described", "sub": [], "rec": False},
            ]},
        ]
    },

    # ── D6c ───────────────────────────────────────────────────────────────────
    {
        "id": "D6c", "title": "PROCEDURE DESCRIPTION — DEVICES & MATERIALS",
        "items": [
            {"id": "6c.1", "text": "All devices used identified (manufacturer, model, size/specifications)", "sub": [], "rec": False},
            {"id": "6c.2", "text": "Embolic agents described when applicable (type, particle size, dose/volume delivered)", "sub": [], "rec": False},
            {"id": "6c.3", "text": "Energy delivery parameters reported when applicable",
             "sub": ["Ablation (RFA, MWA, cryoablation, IRE): energy delivered, duration, target temperature or impedance",
                     "Radioembolization: dosimetry method, prescribed and delivered dose in Gy"],
             "rec": False},
            {"id": "6c.4", "text": "Drug or agent doses reported when applicable (chemotherapy, sclerosant, thrombolytic — concentration, volume, dwell time)", "sub": [], "rec": True},
            {"id": "6c.5", "text": "Device generation or iteration identified when clinically relevant", "sub": [], "rec": True},
        ],
        "conditionals": [
            {"design": "Case Series", "note": "", "items": [
                {"id": "6c.CS.1", "text": "Device or material changes during the study period disclosed", "sub": [], "rec": True},
            ]},
            {"design": "RCT", "note": "", "items": [
                {"id": "6c.RCT.1", "text": "Whether device/material was identical across all sites and operators confirmed", "sub": [], "rec": False},
            ]},
            {"design": "Registry", "note": "", "items": [
                {"id": "6c.REG.1", "text": "Variability in devices and materials across centers described", "sub": [], "rec": False},
                {"id": "6c.REG.2", "text": "Completeness of device/material data fields reported (% available)", "sub": [], "rec": True},
            ]},
        ]
    },

    # ── D7 ────────────────────────────────────────────────────────────────────
    {
        "id": "D7", "title": "PERIPROCEDURAL MANAGEMENT",
        "items": [
            {"id": "7.1", "text": "Anesthesia or sedation type reported",
             "sub": ["If general anesthesia: ASA classification and airway management described",
                     "If conscious sedation: agents, doses, and monitoring level reported",
                     "If local only: explicit statement provided"],
             "rec": False},
            {"id": "7.2", "text": "Antibiotic prophylaxis addressed",
             "sub": ["If used: agent, dose, route, and timing reported",
                     "If not used: explicit statement provided"],
             "rec": False},
            {"id": "7.3", "text": "Antithrombotic management addressed",
             "sub": ["If anticoagulation was held: bridging protocol and resumption timing described",
                     "If anticoagulation was continued: rationale stated",
                     "If reversal agents were used: agent and dose reported"],
             "rec": False},
            {"id": "7.4", "text": "Post-procedural monitoring described",
             "sub": ["If ICU admission: indication and duration reported",
                     "If ward monitoring: duration and parameters reported",
                     "If day-case: discharge criteria explicitly stated"],
             "rec": False},
            {"id": "7.5", "text": "Post-procedural pain management addressed",
             "sub": ["If structured analgesia protocol used: regimen reported",
                     "If post-embolization syndrome anticipated: prophylactic or treatment protocol described"],
             "rec": True},
            {"id": "7.6", "text": "Adjunctive medications administered periprocedurally reported",
             "sub": ["If corticosteroids, antiemetics, vasopressors, or other agents used: drug, dose, and indication described"],
             "rec": True},
        ],
        "conditionals": [
            {"design": "Case Series", "note": "", "items": [
                {"id": "7.CS.1", "text": "Whether a standardized periprocedural protocol was in place — stated explicitly", "sub": [], "rec": False},
            ]},
            {"design": "Cohort with Comparator / RCT", "note": "", "items": [
                {"id": "7.C1", "text": "Periprocedural management confirmed as uniform across groups, or differences explicitly described", "sub": [], "rec": False},
                {"id": "7.C2", "text": "Any adjunctive treatments applied asymmetrically across groups disclosed", "sub": [], "rec": False},
            ]},
            {"design": "RCT", "note": "", "items": [
                {"id": "7.RCT.1", "text": "Co-interventions standardized across sites confirmed, or protocol deviations described", "sub": [], "rec": False},
            ]},
            {"design": "Registry", "note": "", "items": [
                {"id": "7.REG.1", "text": "Variability in periprocedural protocols across centers acknowledged", "sub": [], "rec": True},
            ]},
        ]
    },

    # ── D8 ────────────────────────────────────────────────────────────────────
    {
        "id": "D8", "title": "RADIATION, CONTRAST & DOSE DELIVERY",
        "items": [
            {"id": "8.1", "text": "Radiation exposure metrics reported",
             "sub": ["If fluoroscopy used: fluoroscopy time AND at least one dose metric (DAP, reference point air kerma, or CTDI)",
                     "If CT guidance used: CTDI and DLP reported",
                     "If no ionizing radiation used: explicit statement provided"],
             "rec": False},
            {"id": "8.2", "text": "Radiation protection measures described (shielding, dose reduction protocols, or dose monitoring software)", "sub": [], "rec": True},
            {"id": "8.3", "text": "Contrast agent exposure reported",
             "sub": ["If iodinated contrast used: agent type, concentration, and total volume reported",
                     "If gadolinium used: agent and total dose reported",
                     "If CO₂ or contrast-free: explicit statement provided"],
             "rec": False},
            {"id": "8.4", "text": "Renal function monitoring post-contrast addressed (if pre-existing CKD or high contrast volume: creatinine follow-up timing and results reported)", "sub": [], "rec": True},
            {"id": "8.5", "text": "Cumulative exposure addressed in repeat or staged procedures (cumulative radiation dose and contrast volume reported)", "sub": [], "rec": True},
            {"id": "8.6", "text": "Thermal or energy dose reported for ablative procedures",
             "sub": ["RFA / MWA / cryoablation / IRE: energy delivered, duration, target temperature or impedance",
                     "HIFU: acoustic power and sonication parameters"],
             "rec": False},
            {"id": "8.7", "text": "Radioembolization dosimetry reported (dosimetry method stated; prescribed and delivered dose in Gy reported)", "sub": [], "rec": False},
        ],
        "conditionals": [
            {"design": "Case Series", "note": "", "items": [
                {"id": "8.CS.1", "text": "Variability in exposure metrics across the series described", "sub": [], "rec": True},
            ]},
            {"design": "Cohort with Comparator / RCT", "note": "", "items": [
                {"id": "8.C1", "text": "Radiation and contrast exposure reported separately per group", "sub": [], "rec": False},
                {"id": "8.C2", "text": "Significant exposure differences between groups addressed as potential confounder", "sub": [], "rec": True},
            ]},
            {"design": "Registry", "note": "", "items": [
                {"id": "8.REG.1", "text": "Completeness of radiation and contrast data fields reported (% available)", "sub": [], "rec": True},
            ]},
        ]
    },

    # ── D9 ────────────────────────────────────────────────────────────────────
    {
        "id": "D9", "title": "OUTCOME DEFINITION & TAXONOMY",
        "items": [
            {"id": "9.1",  "text": "Primary endpoint explicitly declared and pre-specified with objective criteria", "sub": [], "rec": False},
            {"id": "9.2",  "text": "Secondary endpoints listed and defined", "sub": [], "rec": False},
            {"id": "9.3",  "text": "Outcomes classified according to the IR taxonomy — technical, imaging/radiological, and clinical outcomes reported separately and not conflated", "sub": [], "rec": False},
            {"id": "9.4",  "text": "Technical success defined using objective, pre-specified criteria",
             "sub": ["If SIR/CIRSE reporting standards exist for the procedure: definition used stated and reference cited",
                     "If no standard exists: definition provided with explicit justification"],
             "rec": False},
            {"id": "9.5",  "text": "Imaging response criteria stated when applicable",
             "sub": ["If oncological procedure: response criteria identified (mRECIST, RECIST 1.1, LI-RADS LR-TR, EASL, WHO — as appropriate)",
                     "If vascular procedure: patency definitions provided (primary, assisted primary, secondary)",
                     "If other: relevant imaging endpoint defined with timepoint"],
             "rec": False},
            {"id": "9.6",  "text": "Clinical success defined separately from technical success with its own objective criteria and assessment timepoint", "sub": [], "rec": False},
            {"id": "9.7",  "text": "Outcome assessment method described",
             "sub": ["If independent adjudication or core lab used: stated explicitly",
                     "If operator-assessed: acknowledged as potential source of bias"],
             "rec": True},
            {"id": "9.8",  "text": "Assessment timepoints pre-specified for all primary and secondary endpoints", "sub": [], "rec": False},
            {"id": "9.9",  "text": "Patient-reported outcomes (PROs) or quality of life instruments described when used (validated instrument identified; administration method and timepoints reported)", "sub": [], "rec": True},
            {"id": "9.10", "text": "Composite endpoints fully disaggregated — all components defined and results of each reported individually", "sub": [], "rec": True},
            {"id": "9.11", "text": "All outcomes reported as both absolute numbers AND proportions (n/N, %)",
             "sub": ["Denominator clearly stated for each timepoint",
                     "Rates must reflect the actual population at risk at each assessment point — not the original enrollment N",
                     "Example: '47 of 52 patients (90.4%)' — not 'success rate was 90.4%'"],
             "rec": True},
        ],
        "conditionals": [
            {"design": "Case Series", "note": "", "items": [
                {"id": "9.CS.1", "text": "Outcomes benchmarked against predefined thresholds or published reference values when no comparator group exists", "sub": [], "rec": True},
            ]},
            {"design": "Cohort with Comparator / RCT", "note": "", "items": [
                {"id": "9.C1", "text": "Primary endpoint identical and assessed by the same method across both groups", "sub": [], "rec": False},
                {"id": "9.C2", "text": "Minimum clinically important difference (MCID) defined for primary endpoint", "sub": [], "rec": True},
            ]},
            {"design": "RCT", "note": "", "items": [
                {"id": "9.RCT.1", "text": "Primary endpoint matches pre-registered endpoint — any changes disclosed with rationale", "sub": [], "rec": False},
                {"id": "9.RCT.2", "text": "Superiority, non-inferiority, or equivalence design declared explicitly", "sub": [], "rec": False},
            ]},
            {"design": "Registry", "note": "", "items": [
                {"id": "9.REG.1", "text": "Outcome ascertainment method described (active follow-up, database linkage, or administrative coding)", "sub": [], "rec": False},
                {"id": "9.REG.2", "text": "Completeness of outcome data reported (% with primary endpoint available)", "sub": [], "rec": False},
            ]},
        ]
    },

    # ── D10 ───────────────────────────────────────────────────────────────────
    {
        "id": "D10", "title": "ADVERSE EVENTS",
        "items": [
            {"id": "10.1",  "text": "Adverse event classification system explicitly stated",
             "sub": ["If SIR classification used: minor (class A/B) and major (class C/D/E/F) reported separately",
                     "If CIRSE classification used: grade stated per event",
                     "If neither: classification system identified and reference cited"],
             "rec": False},
            {"id": "10.1b", "text": "Explicit definitions of major and minor complications provided as used in this study",
             "sub": ["Example: 'Minor complications: SIR class A–B — no therapy or nominal therapy required. Major complications: SIR class C–F — therapy required, unplanned admission, permanent sequelae, or death'",
                     "If procedure-specific definitions used (e.g. post-embolization syndrome grading): criteria explicitly stated",
                     "If composite safety endpoint used: all components individually defined"],
             "rec": False},
            {"id": "10.2",  "text": "All adverse events reported as both absolute numbers AND proportions (n/N, %)",
             "sub": ["Denominator clearly stated per event",
                     "Events reported per procedure AND per patient when these differ"],
             "rec": False},
            {"id": "10.3",  "text": "Causal relationship to procedure stated for each event (procedure-related vs. disease-related vs. indeterminate)", "sub": [], "rec": False},
            {"id": "10.4",  "text": "Timing of adverse events reported (intraprocedural, in-hospital ≤30 days, and post-discharge events reported separately)", "sub": [], "rec": True},
            {"id": "10.5",  "text": "Mortality reported with cause attribution",
             "sub": ["All-cause mortality reported",
                     "Procedure-related mortality reported separately with definition used",
                     "30-day mortality reported as minimum standard"],
             "rec": False},
            {"id": "10.6",  "text": "Method of adverse event capture described",
             "sub": ["If systematic, protocol-driven collection: stated explicitly",
                     "If retrospective chart review only: acknowledged as potential source of underreporting"],
             "rec": False},
            {"id": "10.7",  "text": "Unplanned reintervention or unplanned readmission reported (if occurred: indication, timing, and outcome described)", "sub": [], "rec": True},
        ],
        "conditionals": [
            {"design": "Case Series", "note": "", "items": [
                {"id": "10.CS.1", "text": "Whether adverse event collection was prospective or retrospective — stated explicitly", "sub": [], "rec": False},
            ]},
            {"design": "Cohort with Comparator / RCT", "note": "", "items": [
                {"id": "10.C1", "text": "Adverse events ascertained by the same method and at the same timepoints across groups", "sub": [], "rec": False},
                {"id": "10.C2", "text": "Between-group differences in adverse event rates analyzed and reported", "sub": [], "rec": True},
            ]},
            {"design": "RCT", "note": "", "items": [
                {"id": "10.RCT.1", "text": "Pre-specified safety endpoints reported alongside efficacy endpoints", "sub": [], "rec": False},
                {"id": "10.RCT.2", "text": "Early stopping for safety addressed, or DSMB decisions disclosed", "sub": [], "rec": False},
            ]},
            {"design": "Registry", "note": "", "items": [
                {"id": "10.REG.1", "text": "Adverse event capture method across centers described (mandatory vs. voluntary reporting)", "sub": [], "rec": False},
                {"id": "10.REG.2", "text": "Completeness of adverse event data reported (% with safety data available)", "sub": [], "rec": True},
            ]},
        ]
    },

    # ── D11 ───────────────────────────────────────────────────────────────────
    {
        "id": "D11", "title": "FOLLOW-UP & STATISTICAL ANALYSIS",
        "items": [
            {"id": "11.1", "text": "Follow-up protocol described",
             "sub": ["Clinical follow-up: timing and method stated (in-person, telephone, electronic records)",
                     "Imaging follow-up: modality, timing, and criteria for unscheduled imaging described"],
             "rec": False},
            {"id": "11.2", "text": "Follow-up duration reported (median with range or IQR; minimum follow-up for primary endpoint stated)", "sub": [], "rec": False},
            {"id": "11.3", "text": "Loss to follow-up reported explicitly",
             "sub": ["Number and proportion lost at each timepoint reported (n/N, %)",
                     "Reasons for loss described when available",
                     "Whether loss was random or potentially informative addressed"],
             "rec": False},
            {"id": "11.4", "text": "Reintervention during follow-up reported (planned vs. unplanned distinguished; impact on primary outcome analysis described)", "sub": [], "rec": False},
            {"id": "11.5", "text": "Statistical methods appropriate for the study design described for each endpoint", "sub": [], "rec": False},
            {"id": "11.6", "text": "Descriptive statistics reported with appropriate measures",
             "sub": ["Normally distributed continuous variables: mean ± SD",
                     "Non-normally distributed variables: median with IQR or range",
                     "Categorical variables: n/N and %"],
             "rec": False},
            {"id": "11.7", "text": "Time-to-event analysis described when applicable",
             "sub": ["If Kaplan-Meier used: censoring rules explicitly stated",
                     "If competing risks present (e.g. death before local recurrence): competing risk analysis performed or absence justified"],
             "rec": True},
            {"id": "11.8", "text": "Handling of missing data described",
             "sub": ["If complete case analysis: stated explicitly and proportion of missing data reported",
                     "If imputation used: method described (single vs. multiple imputation)",
                     "If missing data >10% for primary endpoint: sensitivity analysis performed"],
             "rec": True},
            {"id": "11.9", "text": "Statistical software identified (name and version)", "sub": [], "rec": False},
        ],
        "conditionals": [
            {"design": "Case Series", "note": "", "items": [
                {"id": "11.CS.1", "text": "Confidence intervals reported for all primary outcome estimates", "sub": [], "rec": True},
                {"id": "11.CS.2", "text": "Subgroup analyses declared as exploratory and hypothesis-generating only", "sub": [], "rec": True},
            ]},
            {"design": "Cohort with Comparator", "note": "", "items": [
                {"id": "11.CO.1", "text": "Confounding control method described",
                 "sub": ["If multivariable regression: variables included and selection criteria stated",
                         "If propensity score: method described (matching, weighting, or stratification) and balance assessment reported"],
                 "rec": False},
                {"id": "11.CO.2", "text": "Effect measure reported with 95% CI for primary endpoint", "sub": [], "rec": False},
            ]},
            {"design": "RCT", "note": "", "items": [
                {"id": "11.RCT.1", "text": "Sample size calculation reported",
                 "sub": ["Assumptions stated: expected event rate, MCID, alpha, power",
                         "Actual achieved power reported if enrollment fell short of target"],
                 "rec": False},
                {"id": "11.RCT.2", "text": "Analysis population defined (ITT, modified ITT, or per-protocol — declared with rationale)", "sub": [], "rec": False},
                {"id": "11.RCT.3", "text": "Interim analyses and stopping rules disclosed", "sub": [], "rec": True},
                {"id": "11.RCT.4", "text": "Multiplicity corrections applied when multiple primary endpoints or multiple interim analyses performed", "sub": [], "rec": True},
            ]},
            {"design": "Registry", "note": "", "items": [
                {"id": "11.REG.1", "text": "Unit of analysis defined (patient, procedure, or lesion — with justification if procedure or lesion-level)", "sub": [], "rec": False},
                {"id": "11.REG.2", "text": "Clustering by center accounted for in analysis",
                 "sub": ["If multilevel or mixed model used: stated explicitly",
                         "If ignored: acknowledged as limitation"],
                 "rec": True},
            ]},
        ]
    },

    # ── D12 ───────────────────────────────────────────────────────────────────
    {
        "id": "D12", "title": "FUNDING, CONFLICTS OF INTEREST & DATA AVAILABILITY",
        "items": [
            {"id": "12.1", "text": "Funding sources disclosed",
             "sub": ["If industry-funded: sponsor identified and role in study design, data collection, analysis, and manuscript preparation described",
                     "If publicly funded: grant number and funding body identified",
                     "If no external funding: explicit statement provided"],
             "rec": False},
            {"id": "12.2", "text": "Conflicts of interest disclosed for all authors",
             "sub": ["Device or pharmaceutical relationships relevant to the study declared",
                     "Consulting, speaker fees, or equity relevant to devices or agents used disclosed",
                     "If no conflicts: explicit statement provided"],
             "rec": False},
            {"id": "12.3", "text": "Role of sponsor in the study described (whether sponsor had access to data, participated in analysis, or held publication veto — stated explicitly)", "sub": [], "rec": True},
            {"id": "12.4", "text": "Data availability statement provided",
             "sub": ["If data available upon request: contact mechanism stated",
                     "If data deposited in repository: repository name and accession number provided",
                     "If data cannot be shared: reason stated"],
             "rec": True},
            {"id": "12.5", "text": "Statistical analysis code or analysis plan availability addressed (if available: repository or supplementary material location stated)", "sub": [], "rec": True},
        ],
        "conditionals": []
    },
]

# ─── DOCUMENT BUILDER ────────────────────────────────────────────────────────

doc = Document()

# Page setup — A4
section = doc.sections[0]
section.page_width  = Cm(21)
section.page_height = Cm(29.7)
section.left_margin  = Cm(1.8)
section.right_margin = Cm(1.8)
section.top_margin   = Cm(1.8)
section.bottom_margin = Cm(1.5)

# Default style
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(9)

# ─── HEADER BLOCK ────────────────────────────────────────────────────────────

header_table = doc.add_table(rows=1, cols=2)
header_table.style = 'Table Grid'
header_table.autofit = False

col_widths = [Cm(10), Cm(6.9)]
for i, w in enumerate(col_widths):
    for cell in header_table.columns[i].cells:
        cell.width = w

left_cell  = header_table.cell(0, 0)
right_cell = header_table.cell(0, 1)

set_cell_bg(left_cell,  'F0F4FA')
set_cell_bg(right_cell, 'F0F4FA')

left_cell.vertical_alignment = 1  # center

p_title = left_cell.paragraphs[0]
p_title.alignment = WD_ALIGN_PARAGRAPH.LEFT
no_space_before_after(p_title)
r = p_title.add_run('CONCUR-IR')
r.bold = True
r.font.size = Pt(22)
r.font.color.rgb = RGBColor(0x1C, 0x3A, 0x6E)
r.font.name = 'Calibri'

p_sub = left_cell.add_paragraph()
no_space_before_after(p_sub)
p_sub.alignment = WD_ALIGN_PARAGRAPH.LEFT
r2 = p_sub.add_run('Consensus Criteria for Uniform Reporting in Interventional Radiology')
r2.font.size = Pt(9)
r2.font.color.rgb = RGBColor(0x44, 0x60, 0x8A)
r2.font.name = 'Calibri'

right_cell.vertical_alignment = 1
p_info = right_cell.paragraphs[0]
p_info.alignment = WD_ALIGN_PARAGRAPH.RIGHT
no_space_before_after(p_info)
add_run(p_info, f'Draft v0.1  |  {datetime.date.today().year}\n', bold=False, size=8, color=(0x44, 0x60, 0x8A))

p_authors = right_cell.add_paragraph()
no_space_before_after(p_authors)
p_authors.alignment = WD_ALIGN_PARAGRAPH.RIGHT
add_run(p_authors, 'Authors: ________________________________', size=8, color=(0x88, 0x88, 0x88))

doc.add_paragraph()

# ─── INSTRUCTIONS + DESIGN SELECTOR ─────────────────────────────────────────

instr_table = doc.add_table(rows=1, cols=1)
instr_table.style = 'Table Grid'
instr_table.autofit = False
instr_table.columns[0].width = Cm(16.9)
cell = instr_table.cell(0, 0)
set_cell_bg(cell, 'EEF3FA')

p = cell.paragraphs[0]
no_space_before_after(p)
add_run(p, 'Instructions:  ', bold=True, size=8.5, color=(0x1C, 0x3A, 0x6E))
add_run(p, 'Complete all core items for every study. For each applicable study design, expand the corresponding conditional block and complete the additional items. '
           'All items represent best-practice recommendations — no item constitutes a mandatory publication requirement. '
           'Report as both absolute numbers and proportions (n/N, %) wherever applicable.', size=8.5, color=(0x33, 0x33, 0x33))

p2 = cell.add_paragraph()
no_space_before_after(p2)
add_run(p2, 'Select all applicable study designs:  ', bold=True, size=8.5, color=(0x1C, 0x3A, 0x6E))
add_run(p2, '☐  Case Series     ☐  Cohort     ☐  RCT     ☐  Registry     ☐  Case-Control', size=8.5, color=(0x33, 0x33, 0x33))

doc.add_paragraph()

# ─── MAIN CHECKLIST TABLE ─────────────────────────────────────────────────────

# Column widths: ID | Item | Page
COL_W = [Cm(1.5), Cm(13.3), Cm(2.1)]

table = doc.add_table(rows=0, cols=3)
table.style = 'Table Grid'
table.autofit = False
for i, w in enumerate(COL_W):
    for cell in table.columns[i].cells:
        cell.width = w

def add_column_header_row(t):
    row = t.add_row()
    cells = row.cells
    for c in cells:
        set_cell_bg(c, C_PAGE_HEADER)
    data = [('Item', True), ('Reporting item', True), ('Manuscript page No.', True)]
    for i, (txt, bold) in enumerate(data):
        p = cells[i].paragraphs[0]
        no_space_before_after(p)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_run(p, txt, bold=bold, size=8, color=(0x1C, 0x3A, 0x6E))

def add_domain_header_row(t, domain_id, domain_title):
    row = t.add_row()
    merged = row.cells[0]
    merged.merge(row.cells[1])
    merged.merge(row.cells[2])
    set_cell_bg(merged, C_DOMAIN_BG)
    p = merged.paragraphs[0]
    no_space_before_after(p)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    add_run(p, f'  {domain_id}  —  {domain_title}', bold=True, size=9.5,
            color=(0xFF, 0xFF, 0xFF))

def add_item_row(t, item, bg=C_WHITE, indent=''):
    row = t.add_row()
    cells = row.cells
    for c in cells:
        set_cell_bg(c, bg)

    # ID cell
    p_id = cells[0].paragraphs[0]
    no_space_before_after(p_id)
    p_id.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cells[0].vertical_alignment = 1
    add_run(p_id, item['id'], bold=False, size=8, color=(0x44, 0x44, 0x44))

    # Text cell
    p_txt = cells[1].paragraphs[0]
    no_space_before_after(p_txt)

    checkbox = '  ☐  '
    add_run(p_txt, indent + checkbox, bold=False, size=9, color=(0x33, 0x33, 0x33))
    add_run(p_txt, item['text'], bold=False, size=9, color=(0x11, 0x11, 0x11))
    if item.get('rec'):
        add_run(p_txt, '  (recommended)', bold=False, italic=True, size=8,
                color=(0x77, 0x77, 0x77))

    # Sub-items (→)
    for sub in item.get('sub', []):
        p_sub = cells[1].add_paragraph()
        no_space_before_after(p_sub)
        add_run(p_sub, indent + '        → ', bold=False, italic=True, size=8,
                color=(0x44, 0x60, 0x8A))
        add_run(p_sub, sub, bold=False, italic=True, size=8, color=(0x44, 0x44, 0x44))

    # Page cell
    p_pg = cells[2].paragraphs[0]
    no_space_before_after(p_pg)
    p_pg.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cells[2].vertical_alignment = 1

    return row

def add_conditional_header_row(t, design_name):
    row = t.add_row()
    merged = row.cells[0]
    merged.merge(row.cells[1])
    merged.merge(row.cells[2])
    set_cell_bg(merged, C_COND_HEADER)
    p = merged.paragraphs[0]
    no_space_before_after(p)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    add_run(p, f'    ▼  {design_name}', bold=True, size=8.5,
            color=(0xFF, 0xFF, 0xFF))

def add_note_row(t, note_text):
    row = t.add_row()
    merged = row.cells[0]
    merged.merge(row.cells[1])
    merged.merge(row.cells[2])
    set_cell_bg(merged, C_NOTE_BG)
    p = merged.paragraphs[0]
    no_space_before_after(p)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    add_run(p, '    ' + note_text, bold=False, italic=True, size=7.5,
            color=(0x55, 0x55, 0x55))

# ─── RENDER ──────────────────────────────────────────────────────────────────

add_column_header_row(table)

for domain in DOMAINS:
    add_domain_header_row(table, domain['id'], domain['title'])
    for item in domain['items']:
        add_item_row(table, item, bg=C_WHITE)
    for cond in domain['conditionals']:
        add_conditional_header_row(table, cond['design'])
        if cond.get('note'):
            add_note_row(table, cond['note'])
        for item in cond['items']:
            add_item_row(table, item, bg=C_COND_ITEM, indent='    ')

# ─── FOOTER NOTE ─────────────────────────────────────────────────────────────

doc.add_paragraph()
p_footer = doc.add_paragraph()
no_space_before_after(p_footer)
p_footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
add_run(p_footer,
        'CONCUR-IR v0.1  —  Draft for Expert Review  |  '
        'Consensus Criteria for Uniform Reporting in Interventional Radiology  |  '
        f'{datetime.date.today().year}',
        size=7.5, italic=True, color=(0x88, 0x88, 0x88))

# ─── SAVE ────────────────────────────────────────────────────────────────────

out_path = r'c:\Users\JLOW0\OneDrive\Área de Trabalho\IRSET\CONCUR-IR_Checklist_v0.1.docx'
doc.save(out_path)
print(f'Saved: {out_path}')
